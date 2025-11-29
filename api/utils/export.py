"""
诊断报告导出工具类

支持导出格式：
- Markdown (.md)
- JSON (.json)
- Word (.docx)
- PDF (.pdf)
"""
import io
import json
import re
from datetime import datetime
from typing import Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown


class DiagnosisExporter:
    """诊断报告导出器"""

    @staticmethod
    def export_to_pdf(
        diagnosis_markdown: str,
        patient_name: str,
        patient_id: str,
        case_id: int,
        diagnosis_id: Optional[int] = None,
        timestamp: Optional[str] = None
    ) -> bytes:
        """
        导出为PDF格式

        Args:
            diagnosis_markdown: 诊断报告的markdown文本
            patient_name: 患者姓名
            patient_id: 病历号
            case_id: 病例ID
            diagnosis_id: 诊断ID（可选）
            timestamp: 诊断时间（可选）

        Returns:
            PDF文件的字节数据
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # 构建PDF内容
        story = []
        styles = getSampleStyleSheet()

        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=TA_CENTER,
        )

        # 添加标题
        story.append(Paragraph("AI Medical Diagnosis Report", title_style))
        story.append(Spacer(1, 0.2 * inch))

        # 添加患者信息
        info_style = styles['Normal']
        story.append(Paragraph(f"<b>Patient Name:</b> {patient_name}", info_style))
        story.append(Paragraph(f"<b>Patient ID:</b> {patient_id}", info_style))
        story.append(Paragraph(f"<b>Case ID:</b> {case_id}", info_style))
        if diagnosis_id:
            story.append(Paragraph(f"<b>Diagnosis ID:</b> {diagnosis_id}", info_style))
        if timestamp:
            story.append(Paragraph(f"<b>Diagnosis Time:</b> {timestamp}", info_style))

        story.append(Spacer(1, 0.3 * inch))

        # 处理Markdown内容
        lines = diagnosis_markdown.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1 * inch))
                continue

            # 标题处理
            if line.startswith('# '):
                text = line[2:]
                story.append(Paragraph(text, styles['Heading1']))
            elif line.startswith('## '):
                text = line[3:]
                story.append(Paragraph(text, styles['Heading2']))
            elif line.startswith('### '):
                text = line[4:]
                story.append(Paragraph(text, styles['Heading3']))
            # 列表项处理
            elif line.startswith('- ') or line.startswith('* '):
                text = '• ' + line[2:]
                story.append(Paragraph(text, styles['Normal']))
            # 普通段落
            else:
                # 处理加粗 - 使用正则表达式正确处理配对的**
                text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                # 移除其他Markdown符号
                text = text.replace('`', '')
                story.append(Paragraph(text, styles['Normal']))

            story.append(Spacer(1, 0.05 * inch))

        # 生成PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        return pdf_bytes

    @staticmethod
    def export_to_docx(
        diagnosis_markdown: str,
        patient_name: str,
        patient_id: str,
        case_id: int,
        diagnosis_id: Optional[int] = None,
        timestamp: Optional[str] = None
    ) -> bytes:
        """
        导出为Word格式

        Args:
            diagnosis_markdown: 诊断报告的markdown文本
            patient_name: 患者姓名
            patient_id: 病历号
            case_id: 病例ID
            diagnosis_id: 诊断ID（可选）
            timestamp: 诊断时间（可选）

        Returns:
            Word文件的字节数据
        """
        doc = Document()

        # 添加标题
        title = doc.add_heading('AI Medical Diagnosis Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加患者信息
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Patient Name: ').bold = True
        p.add_run(patient_name)

        p = doc.add_paragraph()
        p.add_run('Patient ID: ').bold = True
        p.add_run(patient_id)

        p = doc.add_paragraph()
        p.add_run('Case ID: ').bold = True
        p.add_run(str(case_id))

        if diagnosis_id:
            p = doc.add_paragraph()
            p.add_run('Diagnosis ID: ').bold = True
            p.add_run(str(diagnosis_id))

        if timestamp:
            p = doc.add_paragraph()
            p.add_run('Diagnosis Time: ').bold = True
            p.add_run(timestamp)

        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        doc.add_paragraph()

        # 处理Markdown内容
        lines = diagnosis_markdown.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # 标题处理
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            # 列表项处理
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            # 普通段落
            else:
                p = doc.add_paragraph()
                # 简单处理加粗
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()
        buffer.close()

        return docx_bytes

    @staticmethod
    def export_to_markdown(
        diagnosis_markdown: str,
        patient_name: str,
        patient_id: str,
        case_id: int,
        diagnosis_id: Optional[int] = None,
        timestamp: Optional[str] = None
    ) -> bytes:
        """
        导出为Markdown格式（添加元数据头部）

        Args:
            diagnosis_markdown: 诊断报告的markdown文本
            patient_name: 患者姓名
            patient_id: 病历号
            case_id: 病例ID
            diagnosis_id: 诊断ID（可选）
            timestamp: 诊断时间（可选）

        Returns:
            Markdown文件的字节数据
        """
        content = f"""# AI Medical Diagnosis Report

---

**Patient Name:** {patient_name}
**Patient ID:** {patient_id}
**Case ID:** {case_id}
"""
        if diagnosis_id:
            content += f"**Diagnosis ID:** {diagnosis_id}  \n"
        if timestamp:
            content += f"**Diagnosis Time:** {timestamp}  \n"

        content += f"""
---

{diagnosis_markdown}

---

*Generated by AI Medical Diagnostics System*
"""

        return content.encode('utf-8')

    @staticmethod
    def export_to_json(
        diagnosis_markdown: str,
        patient_name: str,
        patient_id: str,
        case_id: int,
        diagnosis_id: Optional[int] = None,
        timestamp: Optional[str] = None,
        model: Optional[str] = None,
        execution_time_ms: Optional[int] = None
    ) -> bytes:
        """
        导出为JSON格式

        Args:
            diagnosis_markdown: 诊断报告的markdown文本
            patient_name: 患者姓名
            patient_id: 病历号
            case_id: 病例ID
            diagnosis_id: 诊断ID（可选）
            timestamp: 诊断时间（可选）
            model: 使用的模型名称（可选）
            execution_time_ms: 执行时间（毫秒）（可选）

        Returns:
            JSON文件的字节数据
        """
        data = {
            "report_type": "AI Medical Diagnosis Report",
            "patient_info": {
                "name": patient_name,
                "patient_id": patient_id,
            },
            "diagnosis_info": {
                "case_id": case_id,
                "diagnosis_id": diagnosis_id,
                "timestamp": timestamp,
                "model": model,
                "execution_time_ms": execution_time_ms,
            },
            "diagnosis_content": diagnosis_markdown,
            "export_time": datetime.now().isoformat(),
        }

        json_str = json.dumps(data, ensure_ascii=False, indent=2)
        return json_str.encode('utf-8')
